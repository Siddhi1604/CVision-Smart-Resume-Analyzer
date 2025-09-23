from fastapi import FastAPI, UploadFile, File, Form
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Optional, Set, Tuple, Any
import io
import json
import re
import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
import httpx

try:
    import docx  # python-docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    docx = None

# Load environment variables from .env file located next to this file
_BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
_DOTENV_PATH = os.path.join(_BACKEND_DIR, ".env")
load_dotenv(dotenv_path=_DOTENV_PATH, override=True)


app = FastAPI(title="CVision Standard Analyzer", version="0.1.0")

# CORS - allow frontend dev server by default
allowed_origins_env = os.environ.get("BACKEND_ALLOWED_ORIGINS", "*")
allowed_origins = (
    [o.strip() for o in allowed_origins_env.split(",") if o.strip()]
    if allowed_origins_env != "*"
    else ["*"]
)
allow_credentials = allowed_origins != ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=allow_credentials,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OpenAI client for AI analysis
openai_api_key = os.environ.get("OPENROUTER_API_KEY")
if not openai_api_key:
    print("Warning: OPENROUTER_API_KEY not found in environment variables. AI analysis will not work.")
    openai_client = None
else:
    openai_client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=openai_api_key,
        default_headers={
            "HTTP-Referer": "http://localhost:3000",
            "X-Title": "CVision Resume Analyzer",
        }
    )

# Simple in-memory storage for resume analyses (in production, use a database)
resume_analyses_storage = []

# Simple disk persistence
_STORAGE_DIR = os.path.join(_BACKEND_DIR, "storage")
_UPLOADS_DIR = os.path.join(_BACKEND_DIR, "uploads")
_ANALYSES_JSON = os.path.join(_STORAGE_DIR, "analyses.json")
os.makedirs(_STORAGE_DIR, exist_ok=True)
os.makedirs(_UPLOADS_DIR, exist_ok=True)

def _load_analyses_from_disk():
    """Load resume analyses from disk storage"""
    global resume_analyses_storage
    try:
        if os.path.exists(_ANALYSES_JSON):
            with open(_ANALYSES_JSON, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    resume_analyses_storage = data
    except Exception as e:
        print(f"Failed to load analyses from disk: {e}")

def _save_analyses_to_disk():
    """Save resume analyses to disk storage"""
    try:
        with open(_ANALYSES_JSON, "w", encoding="utf-8") as f:
            json.dump(resume_analyses_storage, f, ensure_ascii=False)
    except Exception as e:
        print(f"Failed to save analyses to disk: {e}")

# Initialize storage
_load_analyses_from_disk()

class ResumeAnalysis(BaseModel):
    id: Optional[str] = None
    user_id: str
    resume_name: str
    job_category: str
    job_role: str
    analysis_type: str  # 'standard' or 'ai'
    analysis_result: Dict
    created_at: Optional[str] = None
    file_name: Optional[str] = None


class FeedbackRequest(BaseModel):
    name: str
    email: str
    subject: str
    message: str
    rating: Optional[int] = None


# ==== Resume Builder Request Model ====
class ExperienceItem(BaseModel):
    company: str = ""
    position: str = ""
    startDate: str = ""
    endDate: str = ""
    description: str = ""
    responsibilities: List[str] = []
    achievements: List[str] = []


class EducationItem(BaseModel):
    school: str = ""
    degree: str = ""
    field: str = ""
    graduationDate: str = ""
    gpa: str = ""
    achievements: List[str] = []


class ProjectItem(BaseModel):
    name: str = ""
    technologies: str = ""
    description: str = ""
    responsibilities: List[str] = []
    achievements: List[str] = []
    link: str = ""


class Skills(BaseModel):
    technical: List[str] = []
    soft: List[str] = []
    languages: List[str] = []
    tools: List[str] = []


class PersonalInfo(BaseModel):
    fullName: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    portfolio: str = ""


class ResumeBuildRequest(BaseModel):
    personalInfo: PersonalInfo
    summary: str = ""
    experience: List[ExperienceItem] = []
    education: List[EducationItem] = []
    projects: List[ProjectItem] = []
    skills: Skills = Skills()
    template: str = "Modern"


# ==== Job Search Models ====
class JobSearchRequest(BaseModel):
    page: int = 0
    descending: bool = False
    company: Optional[str] = None
    category: Optional[str] = None
    level: Optional[str] = None
    location: Optional[str] = None


def _add_section_title(d, title: str, template: str):
    p = d.add_paragraph()
    run = p.add_run(title.upper() if template in {"Professional", "Minimal"} else title)
    run.bold = True
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(2)


def _add_bullets(d, items: List[str]):
    for item in items:
        if not item or not item.strip():
            continue
        p = d.add_paragraph(style='List Bullet')
        p.add_run(item.strip())


def _compose_contact_line(pi: PersonalInfo) -> str:
    parts: List[str] = []
    if pi.email:
        parts.append(pi.email)
    if pi.phone:
        parts.append(pi.phone)
    if pi.location:
        parts.append(pi.location)
    if pi.linkedin:
        parts.append(pi.linkedin)
    if pi.portfolio:
        parts.append(pi.portfolio)
    return " \u2022 ".join([p for p in parts if p])


def build_resume_docx(payload: ResumeBuildRequest) -> bytes:
    if docx is None:
        raise HTTPException(status_code=503, detail="Document service not available (python-docx missing)")

    document = docx.Document()

    # Template-specific style setup
    template = (payload.template or "Modern").strip()
    template = template if template in {"Modern", "Professional", "Minimal", "Creative"} else "Modern"

    def apply_color(run, hex_rgb: str):
        if not hex_rgb:
            return
        try:
            hex_rgb = hex_rgb.lstrip('#')
            run.font.color.rgb = RGBColor(int(hex_rgb[0:2], 16), int(hex_rgb[2:4], 16), int(hex_rgb[4:6], 16))
        except Exception:
            pass

    # Header: Name (distinct per template)
    name_paragraph = document.add_paragraph()
    name_run = name_paragraph.add_run(payload.personalInfo.fullName or "Your Name")
    name_run.bold = True

    if template == "Modern":
        name_run.font.size = Pt(22)
        apply_color(name_run, "22c55e")  # green accent
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif template == "Professional":
        name_run.font.size = Pt(18)
        apply_color(name_run, "111827")  # near-black
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif template == "Minimal":
        name_run.font.size = Pt(20)
        apply_color(name_run, "374151")  # gray
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:  # Creative
        name_run.font.size = Pt(24)
        apply_color(name_run, "a78bfa")  # purple accent
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Optional accent divider for Modern/Creative
    if template in {"Modern", "Creative"}:
        divider = document.add_paragraph()
        dr = divider.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
        apply_color(dr, "e5e7eb")
        divider.paragraph_format.space_after = Pt(4)

    # Contact line (distinct alignment + subtle color)
    contact = _compose_contact_line(payload.personalInfo)
    if contact:
        contact_p = document.add_paragraph(contact)
        if template in {"Modern", "Creative"}:
            contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            contact_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        contact_p.paragraph_format.space_after = Pt(10 if template == "Minimal" else 8)

    # Summary
    if payload.summary and payload.summary.strip():
        _add_section_title(document, "Summary", template)
        p = document.add_paragraph(payload.summary.strip())
        if template == "Minimal":
            p.paragraph_format.space_after = Pt(12)

    # Skills
    if any([payload.skills.technical, payload.skills.soft, payload.skills.languages, payload.skills.tools]):
        _add_section_title(document, "Skills", template)
        def add_skill_line(label: str, values: List[str]):
            if values:
                p = document.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(", ".join(values))
        add_skill_line("Technical", payload.skills.technical)
        add_skill_line("Soft", payload.skills.soft)
        add_skill_line("Languages", payload.skills.languages)
        add_skill_line("Tools", payload.skills.tools)

    # Experience
    if payload.experience:
        _add_section_title(document, "Experience", template)
        for exp in payload.experience:
            header = ", ".join([p for p in [exp.position, exp.company] if p])
            dates = " - ".join([p for p in [exp.startDate, exp.endDate] if p])
            title_line = header + (f"  |  {dates}" if dates else "")
            if title_line:
                p = document.add_paragraph()
                r = p.add_run(title_line)
                r.bold = True
                if template == "Professional":
                    apply_color(r, "1f2937")
            if exp.description and exp.description.strip():
                document.add_paragraph(exp.description.strip())
            _add_bullets(document, exp.responsibilities or [])
            _add_bullets(document, exp.achievements or [])

    # Education
    if payload.education:
        _add_section_title(document, "Education", template)
        for edu in payload.education:
            header = ", ".join([p for p in [edu.degree, edu.field] if p])
            school_part = edu.school or ""
            date_part = edu.graduationDate or ""
            line_parts = [part for part in [header, school_part, date_part] if part]
            if line_parts:
                p = document.add_paragraph()
                rr = p.add_run(" | ".join(line_parts))
                rr.bold = True
                if template == "Minimal":
                    apply_color(rr, "4b5563")
            if edu.gpa:
                document.add_paragraph(f"GPA: {edu.gpa}")
            _add_bullets(document, edu.achievements or [])

    # Projects
    if payload.projects:
        _add_section_title(document, "Projects", template)
        for proj in payload.projects:
            header = proj.name or ""
            tech = proj.technologies or ""
            if header:
                p = document.add_paragraph()
                run = p.add_run(header)
                run.bold = True
                if tech:
                    p.add_run(f" — {tech}")
            if proj.description:
                document.add_paragraph(proj.description)
            _add_bullets(document, proj.responsibilities or [])
            _add_bullets(document, proj.achievements or [])
            if proj.link:
                document.add_paragraph(proj.link)

    # Serialize to bytes
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


@app.post("/build-resume")
async def build_resume(req: ResumeBuildRequest):
    try:
        content = build_resume_docx(req)
        filename = (req.personalInfo.fullName or "resume").replace(" ", "_") + ".docx"
        return StreamingResponse(
            io.BytesIO(content),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=\"{filename}\""
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error building resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to build resume")


def send_feedback_email(feedback: FeedbackRequest):
    """Send feedback email to the team"""
    try:
        # Email configuration
        sender_email = "cvision.feedback@gmail.com"  # You'll need to set up this email
        sender_password = os.environ.get("EMAIL_PASSWORD", "")  # Set this in .env
        recipients = ["22it084@charusat.edu.in", "22it157@charusat.edu.in"]
        
        # Create message
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = ", ".join(recipients)
        msg['Subject'] = f"CVision Feedback: {feedback.subject}"
        
        # Create email body
        body = f"""
        New feedback received for CVision Resume Analyzer
        
        From: {feedback.name} ({feedback.email})
        Subject: {feedback.subject}
        Rating: {feedback.rating}/5 (if provided)
        Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        
        Message:
        {feedback.message}
        
        ---
        This feedback was sent automatically from the CVision application.
        """
        
        msg.attach(MIMEText(body, 'plain'))
        
        # Send email using Gmail SMTP
        if sender_password:
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            server.quit()
            return True
        else:
            print("Email password not configured. Feedback would be sent to:", recipients)
            print("Feedback content:", body)
            return False
            
    except Exception as e:
        print(f"Error sending feedback email: {e}")
        return False


def _load_roles_dataset() -> Dict[str, Dict[str, List[str]]]:
    """Load job roles dataset from roles.json file"""
    try:
        repo_root = os.path.dirname(os.path.abspath(__file__))
        roles_path = os.path.join(repo_root, "roles.json")
        with open(roles_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


ROLES_DATASET = _load_roles_dataset()

# ==== Job Search Service ====


@app.get("/job-categories")
async def list_job_categories():
    return {"categories": sorted(list(ROLES_DATASET.keys()))}


@app.get("/job-roles")
async def list_job_roles(category: Optional[str] = None):
    if not category:
        # Return nested structure compatible with frontend fallback shape
        # { category: { role: { description: str, required_skills: [str] } } }
        nested: Dict[str, Dict[str, Dict[str, List[str]]]] = {}
        for cat, roles in ROLES_DATASET.items():
            nested[cat] = {}
            for role, skills in roles.items():
                nested[cat][role] = {
                    "description": "",
                    "required_skills": skills,
                }
        return nested
    roles = ROLES_DATASET.get(category, {})
    return {"category": category, "roles": sorted(list(roles.keys()))}


def extract_text_from_upload(file: UploadFile) -> str:
    """Extract text content from uploaded resume file (PDF, DOCX, or plain text)"""
    data = file.file.read()
    # Reset pointer so other functions can re-read if needed
    try:
        file.file.seek(0)
    except Exception:
        pass

    filename = (file.filename or "").lower()
    if filename.endswith(".pdf"):
        try:
            from pdfminer.high_level import extract_text as pdf_text
            return pdf_text(io.BytesIO(data)) or ""
        except Exception:
            return ""
    elif filename.endswith(".docx"):
        try:
            import docx  # python-docx
            doc = docx.Document(io.BytesIO(data))
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception:
            return ""
    else:
        # Fallback plain text
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def word_present(text: str, term: str) -> bool:
    """Check if a word is present in text using word boundaries"""
    return re.search(rf"\b{re.escape(term)}\b", text, flags=re.I) is not None

def normalize_text(text: str) -> str:
    """Normalize text for analysis by cleaning and standardizing format"""
    lowered = text.lower()
    # Replace punctuation and slashes with spaces while preserving tokens like c++ -> c++
    cleaned = re.sub(r"[^a-z0-9+#./-]", " ", lowered)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned

def tokenize(text: str) -> Tuple[List[str], List[str]]:
    """Tokenize text into words and bigrams for analysis"""
    tokens = re.findall(r"[a-z0-9+.#-]+", text.lower())
    bigrams = [f"{tokens[i]} {tokens[i+1]}" for i in range(len(tokens)-1)] if len(tokens) > 1 else []
    return tokens, bigrams


ALIASES: Dict[str, List[str]] = {
    # Languages / Core
    "javascript": ["js", "nodejs", "node.js", "vanilla js", "ecmascript"],
    "typescript": ["ts"],
    "python": ["py"],
    "java": [],
    "sql": ["postgres", "postgresql", "mysql", "sqlite"],
    "nosql": ["mongodb", "dynamo", "dynamodb", "cassandra"],
    "html": ["html5"],
    "css": ["css3", "scss", "sass", "tailwind", "bootstrap"],

    # Frameworks / Libraries
    "react": ["reactjs", "react.js", "next", "nextjs", "next.js"],
    "next.js": ["next", "nextjs"],
    "redux": ["redux toolkit", "rtk"],
    "testing-library": ["react testing library"],
    "cypress": ["e2e", "end to end testing"],
    "webpack": [],
    "vite": [],
    "express": ["expressjs", "express.js"],
    "fastapi": [],
    "spring": ["spring boot", "springboot"],
    "graphql": [],

    # Infra / DevOps
    "docker": [],
    "kubernetes": ["k8s"],
    "aws": ["amazon web services", "ec2", "s3", "lambda", "rds"],
    "ci/cd": ["cicd", "continuous integration", "continuous delivery", "github actions", "gitlab ci"],
    "git": ["github", "gitlab", "bitbucket"],
    "redis": [],
    "caching": ["cache"],

    # Concepts
    "rest": ["restful", "rest api", "apis"],
    "system design": ["architecture", "scalability"],
    "design patterns": ["patterns"],
    "algorithms": ["algo"],
    "data structures": ["ds"],
    "machine learning": ["ml", "mlops"],
    "model evaluation": ["evaluation", "metrics"],
    "feature engineering": ["features"],
}


def tokens_contain(tokens: List[str], target: str) -> bool:
    """Check if target string is contained in tokens with fuzzy matching"""
    t = target.lower()
    if t in tokens:
        return True
    # Accept token variants with punctuation removed
    t_compact = re.sub(r"[^a-z0-9]+", "", t)
    for tok in tokens:
        if tok == t:
            return True
        if re.sub(r"[^a-z0-9]+", "", tok) == t_compact:
            return True
    return False

def fuzzy_similar(a: str, b: str, threshold: float = 0.9) -> bool:
    """Check if two strings are similar using sequence matching"""
    from difflib import SequenceMatcher
    return SequenceMatcher(None, a.lower(), b.lower()).ratio() >= threshold

def match_skill(resume_tokens: List[str], resume_bigrams: List[str], skill: str) -> bool:
    """Match a skill against resume tokens using aliases and fuzzy matching"""
    candidates = [skill] + ALIASES.get(skill.lower(), [])
    # Exact token or bigram match
    for c in candidates:
        if tokens_contain(resume_tokens, c) or tokens_contain(resume_bigrams, c):
            return True
    # Substring within tokens (e.g., "typescript" in "ts/tsx/typescript")
    for c in candidates:
        c_compact = re.sub(r"[^a-z0-9]+", "", c.lower())
        for tok in resume_tokens + resume_bigrams:
            tok_compact = re.sub(r"[^a-z0-9]+", "", tok.lower())
            if c_compact and tok_compact and c_compact in tok_compact:
                return True
            if fuzzy_similar(tok_compact, c_compact, threshold=0.92) and len(c_compact) >= 4:
                return True
    return False

def score_keyword_match(text: str, skills: List[str]):
    """Score how well resume matches required skills"""
    normalized = normalize_text(text)
    tokens, bigrams = tokenize(normalized)
    present = []
    for s in skills:
        if match_skill(tokens, bigrams, s):
            present.append(s)
    missing = [s for s in skills if s not in present]
    score = round((len(present) / max(1, len(skills))) * 100)
    return score, missing

def score_sections(text: str) -> int:
    """Score resume based on presence of standard sections"""
    sections = [
        "summary", "objective", "skills", "experience", "employment",
        "education", "projects", "certifications", "contact"
    ]
    found = sum(1 for s in sections if word_present(text, s))
    return round((found / len(sections)) * 100)

def get_present_sections(text: str) -> List[str]:
    """Return list of section names detected in the resume text"""
    sections = [
        "summary", "objective", "skills", "experience", "employment",
        "education", "projects", "certifications", "contact"
    ]
    return [s for s in sections if word_present(text, s)]

def score_format(text: str, raw_len: int) -> int:
    """Score resume format and structure"""
    score = 100
    words = len(re.findall(r"\w+", text))
    if words < 250:
        score -= 20
    if words > 3000:
        score -= 15
    if "•" not in text and "-" not in text:
        score -= 10
    if text.strip() == "" or (raw_len > 0 and len(text) < 50):
        score -= 40
    return max(0, min(100, score))


class AnalyzeResponse(BaseModel):
    ats_score: int
    keyword_match: Dict[str, int]
    missing_skills: List[str]
    format_score: int
    section_score: int
    suggestions: List[str]
    jd_match_score: Optional[int] = None
    contact: Dict[str, bool]
    metrics: Dict[str, int]


@app.post("/analyze-resume", response_model=AnalyzeResponse)
async def analyze_resume(
    file: Optional[UploadFile] = File(None),
    job_category: str = Form(...),
    job_role: str = Form(...),
    custom_job_description: Optional[str] = Form(None),
    text: Optional[str] = Form(None),
    user_id: str = Form("default_user"),
):
    if not file and not text:
        raise HTTPException(status_code=400, detail="Provide either a file or text")

    raw = b""
    if file:
        raw = file.file.read()
        try:
            file.file.seek(0)
        except Exception:
            pass

    resume_text = (text or "").strip() or (extract_text_from_upload(file) if file else "")

    skills = ROLES_DATASET.get(job_category, {}).get(job_role, [])
    km_score, missing = score_keyword_match(resume_text, skills)
    sec_score = score_sections(resume_text)
    fmt_score = score_format(resume_text, len(raw))
    ats = round(0.5 * km_score + 0.25 * sec_score + 0.25 * fmt_score)

    suggestions: List[str] = []
    if missing:
        suggestions.append(
            "Include relevant skills if applicable: " + ", ".join(missing[:8]) + "."
        )
    if km_score < 70:
        suggestions.append(
            "Add more role-specific keywords across Skills and Experience sections."
        )
    if sec_score < 70:
        suggestions.append(
            "Ensure key sections like Summary, Skills, Experience, and Education are present and clearly labeled."
        )
    if fmt_score < 80:
        suggestions.append(
            "Use bullet points and ensure the document text is selectable (avoid image-only PDFs)."
        )
    if custom_job_description:
        suggestions.append(
            "Tailor quantified achievements to the provided job description."
        )

    # Contact info checks
    contact = {
        "has_email": bool(re.search(r"[\w.+'-]+@[\w.-]+\.[A-Za-z]{2,}", resume_text)),
        "has_phone": bool(re.search(r"(\+?\d[\s-]?){7,}\d", resume_text)),
        "has_linkedin": "linkedin.com" in resume_text.lower(),
        "has_github": "github.com" in resume_text.lower(),
    }
    if not contact["has_email"]:
        suggestions.append("Add a professional email address in the header.")
    if not contact["has_phone"]:
        suggestions.append("Include a reachable phone number.")
    if not (contact["has_linkedin"] or contact["has_github"]):
        suggestions.append("Add a LinkedIn or GitHub link if relevant.")

    # JD match score (lightweight)
    jd_match_score: Optional[int] = None
    if custom_job_description:
        def tokenize(s: str) -> Set[str]:
            tokens = re.findall(r"[A-Za-z][A-Za-z0-9+.#-]{1,}", s.lower())
            stop = {
                "and","or","the","a","an","to","for","with","of","in","on","by","at","from","as","is","are","be","this","that","these","those","you","we","they","it","your","our"}
            return {t for t in tokens if len(t) > 2 and t not in stop}

        jd_tokens = tokenize(custom_job_description)
        res_tokens = tokenize(resume_text)
        if jd_tokens:
            overlap = len(jd_tokens & res_tokens)
            jd_match_score = round((overlap / len(jd_tokens)) * 100)
            if jd_match_score < 50:
                suggestions.append("Mirror the language of the job description where appropriate.")

    # Metrics
    word_count = len(re.findall(r"\w+", resume_text))
    reading_time_minutes = max(1, round(word_count / 200))

    result = {
        "ats_score": ats,
        "keyword_match": {"score": km_score},
        "missing_skills": missing,
        "format_score": fmt_score,
        "section_score": sec_score,
        "suggestions": suggestions,
        "jd_match_score": jd_match_score,
        "contact": contact,
        "metrics": {"word_count": word_count, "reading_time_minutes": reading_time_minutes},
    }
    
    # Store the analysis for dashboard
    try:
        import uuid
        from datetime import datetime
        
        # Persist upload to disk if provided
        saved_path = None
        if file and raw:
            safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", file.filename or "resume")
            unique_prefix = datetime.now().strftime("%Y%m%d%H%M%S%f")
            saved_path = os.path.join(_UPLOADS_DIR, f"{unique_prefix}_{safe_name}")
            try:
                with open(saved_path, "wb") as out:
                    out.write(raw)
            except Exception as e:
                print(f"Failed to save upload: {e}")

        analysis_data = {
            "id": str(uuid.uuid4()),
            "user_id": user_id or "default_user",
            "resume_name": file.filename if file else "Text Resume",
            "job_category": job_category,
            "job_role": job_role,
            "analysis_type": "standard",
            "analysis_result": result,
            "created_at": datetime.now().isoformat(),
            "file_name": file.filename if file else None,
            "file_path": saved_path,
            "file_mime": (
                "application/pdf" if (file and (file.filename or "").lower().endswith(".pdf")) else (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if (file and (file.filename or "").lower().endswith(".docx")) else "text/plain"
                )
            ) if file else None,
        }
        resume_analyses_storage.append(analysis_data)
        _save_analyses_to_disk()
    except Exception as e:
        print(f"Failed to store analysis: {e}")
    
    return result


@app.post("/ai-analyze-resume", response_model=AnalyzeResponse)
async def ai_analyze_resume(
    file: Optional[UploadFile] = File(None),
    job_category: str = Form(...),
    job_role: str = Form(...),
    custom_job_description: Optional[str] = Form(None),
    text: Optional[str] = Form(None),
    user_id: str = Form("default_user"),
):
    if not file and not text:
        raise HTTPException(status_code=400, detail="Provide either a file or text")

    raw = b""
    if file:
        raw = file.file.read()
        try:
            file.file.seek(0)
        except Exception:
            pass

    resume_text = (text or "").strip() or (extract_text_from_upload(file) if file else "")
    
    if not resume_text.strip():
        raise HTTPException(status_code=400, detail="No text could be extracted from the provided file")

    # Get role skills for context
    skills = ROLES_DATASET.get(job_category, {}).get(job_role, [])
    present_sections = get_present_sections(resume_text)
    
    # Build the AI prompt (broader, avoids redundant suggestions)
    prompt = f"""You are an expert resume analyst and career coach.
Analyze the following resume for a {job_role} position within {job_category}.
Read the entire resume holistically (not only the skills list). Infer synonyms and equivalents.
Only suggest adding sections or items if they are genuinely missing. The following sections were detected: {present_sections}.
Avoid recommending items already present (e.g., certifications, experience) unless the improvement is about quality/clarity.
Consider both technical and non-technical aspects: impact, outcomes, leadership, collaboration, communication, quantification, clarity, readability, formatting consistency, and relevance.
Tailor feedback to the target role and the provided job description if present.

Resume Text (trimmed):
{resume_text[:3000]}

Target Role: {job_role}
Target Category: {job_category}
Required Skills: {', '.join(skills)}
{f'Job Description: {custom_job_description}' if custom_job_description else ''}

Return ONLY valid JSON in this schema:
{{
  "ats_score": 0-100,
  "keyword_match": {{"score": 0-100}},
  "missing_skills": ["..."],
  "format_score": 0-100,
  "section_score": 0-100,
  "suggestions": ["Concrete, deduplicated, non-generic improvements that do not ask to add things already present"],
  "jd_match_score": 0-100 or null,
  "contact": {{"has_email": bool, "has_phone": bool, "has_linkedin": bool, "has_github": bool}},
  "metrics": {{"word_count": number, "reading_time_minutes": number}}
}}
"""

    if not openai_client:
        raise HTTPException(status_code=503, detail="AI analysis service not configured. Please set OPENROUTER_API_KEY environment variable.")
    
    try:
        completion = openai_client.chat.completions.create(
            model="openai/gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=1500,
        )
        ai_response = completion.choices[0].message.content.strip()
        # Try to parse the JSON response
        try:
            if ai_response.startswith("```json"):
                ai_response = ai_response[7:]
            if ai_response.endswith("```"):
                ai_response = ai_response[:-3]
            analysis_data = json.loads(ai_response)
            result = {
                "ats_score": analysis_data.get("ats_score", 0),
                "keyword_match": analysis_data.get("keyword_match", {"score": 0}),
                "missing_skills": analysis_data.get("missing_skills", []),
                "format_score": analysis_data.get("format_score", 0),
                "section_score": analysis_data.get("section_score", 0),
                "suggestions": analysis_data.get("suggestions", []),
                "jd_match_score": analysis_data.get("jd_match_score"),
                "contact": analysis_data.get("contact", {"has_email": False, "has_phone": False, "has_linkedin": False, "has_github": False}),
                "metrics": analysis_data.get("metrics", {"word_count": len(re.findall(r"\\w+", resume_text)), "reading_time_minutes": max(1, round(len(re.findall(r"\\w+", resume_text)) / 200))}),
            }
            
            # Store the analysis for dashboard
            try:
                import uuid
                from datetime import datetime
            
                # Persist upload to disk if provided
                saved_path = None
                if file and raw:
                    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", file.filename or "resume")
                    unique_prefix = datetime.now().strftime("%Y%m%d%H%M%S%f")
                    saved_path = os.path.join(_UPLOADS_DIR, f"{unique_prefix}_{safe_name}")
                    try:
                        with open(saved_path, "wb") as out:
                            out.write(raw)
                    except Exception as e:
                        print(f"Failed to save upload: {e}")

                analysis_data = {
                    "id": str(uuid.uuid4()),
                    "user_id": user_id or "default_user",
                    "resume_name": file.filename if file else "Text Resume",
                    "job_category": job_category,
                    "job_role": job_role,
                    "analysis_type": "ai",
                    "analysis_result": result,
                    "created_at": datetime.now().isoformat(),
                    "file_name": file.filename if file else None,
                    "file_path": saved_path,
                    "file_mime": (
                        "application/pdf" if (file and (file.filename or "").lower().endswith(".pdf")) else (
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if (file and (file.filename or "").lower().endswith(".docx")) else "text/plain"
                        )
                    ) if file else None,
                }
                resume_analyses_storage.append(analysis_data)
                _save_analyses_to_disk()
            except Exception as e:
                print(f"Failed to store AI analysis: {e}")
            
            return result
        except json.JSONDecodeError as e:
            print(f"AI response parsing failed: {e}")
            print(f"AI response: {ai_response}")
            raise HTTPException(status_code=500, detail="AI analysis failed - using standard analysis")
    except Exception as e:
        print(f"OpenAI API error: {e}")
        raise HTTPException(status_code=500, detail="AI analysis service unavailable")


@app.get("/health")
async def health():
    return {"status": "ok"}


@app.get("/job-skills")
async def list_job_skills(category: str, role: str):
    skills = ROLES_DATASET.get(category, {}).get(role)
    if skills is None:
        raise HTTPException(status_code=404, detail="Category or role not found")
    return {"category": category, "role": role, "skills": skills}


@app.post("/store-analysis")
async def store_analysis(analysis: ResumeAnalysis):
    import uuid
    from datetime import datetime
    
    analysis.id = str(uuid.uuid4())
    analysis.created_at = datetime.now().isoformat()
    
    resume_analyses_storage.append(analysis.dict())
    return {"id": analysis.id, "message": "Analysis stored successfully"}


@app.get("/user-analyses/{user_id}")
async def get_user_analyses(user_id: str):
    user_analyses = [a for a in resume_analyses_storage if a["user_id"] == user_id]
    return {"analyses": user_analyses}
@app.get("/download-resume/{analysis_id}")
async def download_resume(analysis_id: str):
    try:
        record = next((a for a in resume_analyses_storage if a["id"] == analysis_id), None)
        if not record:
            raise HTTPException(status_code=404, detail="Analysis not found")
        file_path = record.get("file_path")
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="No original file stored for this analysis")

        filename = record.get("file_name") or os.path.basename(file_path)
        mime = record.get("file_mime") or "application/octet-stream"
        return StreamingResponse(
            open(file_path, "rb"),
            media_type=mime,
            headers={
                "Content-Disposition": f"attachment; filename=\"{filename}\""
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"Failed to download resume: {e}")
        raise HTTPException(status_code=500, detail="Failed to download resume")


@app.post("/send-feedback")
async def send_feedback(feedback: FeedbackRequest):
    """Send feedback email to the team"""
    try:
        success = send_feedback_email(feedback)
        if success:
            return {"message": "Feedback sent successfully!", "status": "success"}
        else:
            return {"message": "Feedback logged (email not configured)", "status": "logged"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send feedback: {str(e)}")


# ==== Job Data Functions ====

def get_enhanced_mock_jobs(page: int = 0, keyword: str = "", location: str = "") -> dict:
    """Enhanced mock data as fallback"""
    all_jobs = [
        {
            "id": 1,
            "title": "Senior Software Engineer",
            "company": "GitHub",
            "location": "San Francisco, CA",
            "type": "Full-time",
            "salary": "$120,000 - $180,000",
            "experience": "Senior Level",
            "description": "Join GitHub's engineering team to build the future of software development. Work on distributed systems, API design, and developer tools used by millions worldwide.",
            "posted": "2 days ago",
            "logo": "GH",
            "landing_page": "https://github.com/careers",
            "categories": ["Engineering"],
            "tags": ["Python", "Go", "Kubernetes", "Cloud"]
        },
        {
            "id": 2,
            "title": "Frontend Developer",
            "company": "Microsoft",
            "location": "Seattle, WA",
            "type": "Full-time",
            "salary": "$100,000 - $150,000",
            "experience": "Mid Level",
            "description": "Build user-facing features for Microsoft's cloud platforms. Collaborate with design and backend teams to create intuitive user experiences.",
            "posted": "1 week ago",
            "logo": "MS",
            "landing_page": "https://careers.microsoft.com",
            "categories": ["Engineering"],
            "tags": ["React", "TypeScript", "Azure", "CSS"]
        },
        {
            "id": 3,
            "title": "DevOps Engineer",
            "company": "Amazon",
            "location": "Seattle, WA",
            "type": "Full-time",
            "salary": "$110,000 - $160,000",
            "experience": "Mid Level",
            "description": "Manage infrastructure and deployment pipelines for AWS services. Automate processes and ensure high availability of critical systems.",
            "posted": "3 days ago",
            "logo": "AM",
            "landing_page": "https://amazon.jobs",
            "categories": ["Engineering"],
            "tags": ["AWS", "Docker", "Kubernetes", "Terraform"]
        },
        {
            "id": 4,
            "title": "Data Scientist",
            "company": "Google",
            "location": "Mountain View, CA",
            "type": "Full-time",
            "salary": "$130,000 - $200,000",
            "experience": "Senior Level",
            "description": "Apply machine learning and data analysis to solve complex problems. Work with large datasets to drive product decisions and user insights.",
            "posted": "5 days ago",
            "logo": "GO",
            "landing_page": "https://careers.google.com",
            "categories": ["Data Science"],
            "tags": ["Python", "Machine Learning", "TensorFlow", "SQL"]
        },
        {
            "id": 5,
            "title": "Full Stack Developer",
            "company": "Netflix",
            "location": "Los Gatos, CA",
            "type": "Full-time",
            "salary": "$115,000 - $170,000",
            "experience": "Mid Level",
            "description": "Develop end-to-end solutions for Netflix's streaming platform. Build features that enhance the viewing experience for 200+ million subscribers.",
            "posted": "1 week ago",
            "logo": "NF",
            "landing_page": "https://jobs.netflix.com",
            "categories": ["Engineering"],
            "tags": ["JavaScript", "Node.js", "React", "AWS"]
        },
        {
            "id": 6,
            "title": "Backend Engineer",
            "company": "Stripe",
            "location": "San Francisco, CA",
            "type": "Full-time",
            "salary": "$125,000 - $185,000",
            "experience": "Senior Level",
            "description": "Build scalable payment processing systems. Design APIs and services that handle billions of dollars in transactions securely and reliably.",
            "posted": "4 days ago",
            "logo": "ST",
            "landing_page": "https://stripe.com/jobs",
            "categories": ["Engineering"],
            "tags": ["Python", "PostgreSQL", "Redis", "Microservices"]
        },
        {
            "id": 7,
            "title": "Mobile Developer",
            "company": "Uber",
            "location": "San Francisco, CA",
            "type": "Full-time",
            "salary": "$105,000 - $155,000",
            "experience": "Mid Level",
            "description": "Develop mobile applications for Uber's platform. Create seamless experiences for riders and drivers across iOS and Android platforms.",
            "posted": "6 days ago",
            "logo": "UB",
            "landing_page": "https://www.uber.com/careers",
            "categories": ["Engineering"],
            "tags": ["React Native", "iOS", "Android", "JavaScript"]
        },
        {
            "id": 8,
            "title": "Security Engineer",
            "company": "Cloudflare",
            "location": "San Francisco, CA",
            "type": "Full-time",
            "salary": "$120,000 - $180,000",
            "experience": "Senior Level",
            "description": "Protect Cloudflare's infrastructure and customer data. Implement security measures and respond to threats across our global network.",
            "posted": "2 days ago",
            "logo": "CF",
            "landing_page": "https://www.cloudflare.com/careers",
            "categories": ["Security"],
            "tags": ["Security", "Python", "Go", "Network Security"]
        },
        {
            "id": 9,
            "title": "Product Manager",
            "company": "Meta",
            "location": "Menlo Park, CA",
            "type": "Full-time",
            "salary": "$140,000 - $200,000",
            "experience": "Senior Level",
            "description": "Lead product strategy for Meta's family of apps. Work with engineering and design teams to build products used by billions of people.",
            "posted": "1 day ago",
            "logo": "ME",
            "landing_page": "https://www.metacareers.com",
            "categories": ["Product"],
            "tags": ["Product Management", "Strategy", "Analytics", "Leadership"]
        },
        {
            "id": 10,
            "title": "Junior Software Engineer",
            "company": "Slack",
            "location": "Remote",
            "type": "Full-time",
            "salary": "$80,000 - $120,000",
            "experience": "Entry Level",
            "description": "Join Slack's engineering team as a junior developer. Learn from experienced engineers while building features for workplace collaboration.",
            "posted": "3 days ago",
            "logo": "SL",
            "landing_page": "https://slack.com/careers",
            "categories": ["Engineering"],
            "tags": ["JavaScript", "Python", "API", "Collaboration Tools"]
        },
        {
            "id": 11,
            "title": "UX Designer",
            "company": "Airbnb",
            "location": "San Francisco, CA",
            "type": "Full-time",
            "salary": "$110,000 - $160,000",
            "experience": "Mid Level",
            "description": "Design intuitive experiences for Airbnb's platform. Research user needs and create designs that help people belong anywhere.",
            "posted": "1 week ago",
            "logo": "AB",
            "landing_page": "https://careers.airbnb.com",
            "categories": ["Design"],
            "tags": ["UI/UX", "Figma", "User Research", "Design Systems"]
        },
        {
            "id": 12,
            "title": "Machine Learning Engineer",
            "company": "OpenAI",
            "location": "San Francisco, CA",
            "type": "Full-time",
            "salary": "$150,000 - $250,000",
            "experience": "Senior Level",
            "description": "Develop and deploy large-scale machine learning models. Work on cutting-edge AI research and applications that benefit humanity.",
            "posted": "Today",
            "logo": "OA",
            "landing_page": "https://openai.com/careers",
            "categories": ["AI/ML"],
            "tags": ["Machine Learning", "PyTorch", "Deep Learning", "NLP"]
        },
        {
            "id": 13,
            "title": "React Developer",
            "company": "Spotify",
            "location": "Stockholm, Sweden",
            "type": "Full-time",
            "salary": "$95,000 - $140,000",
            "experience": "Mid Level",
            "description": "Build engaging user interfaces for Spotify's music streaming platform. Work with React, TypeScript, and modern web technologies.",
            "posted": "2 days ago",
            "logo": "SP",
            "landing_page": "https://www.spotifyjobs.com",
            "categories": ["Engineering"],
            "tags": ["React", "TypeScript", "JavaScript", "Frontend"]
        },
        {
            "id": 14,
            "title": "Python Developer",
            "company": "Dropbox",
            "location": "Remote",
            "type": "Full-time",
            "salary": "$110,000 - $165,000",
            "experience": "Mid Level",
            "description": "Develop backend services and APIs using Python. Work on distributed systems that handle billions of files.",
            "posted": "1 week ago",
            "logo": "DB",
            "landing_page": "https://jobs.dropbox.com",
            "categories": ["Engineering"],
            "tags": ["Python", "Django", "API", "Distributed Systems"]
        },
        {
            "id": 15,
            "title": "DevOps Specialist",
            "company": "GitLab",
            "location": "Remote",
            "type": "Full-time",
            "salary": "$100,000 - $150,000",
            "experience": "Senior Level",
            "description": "Manage CI/CD pipelines and infrastructure for GitLab's platform. Work with Kubernetes, Docker, and cloud technologies.",
            "posted": "3 days ago",
            "logo": "GL",
            "landing_page": "https://about.gitlab.com/jobs",
            "categories": ["Engineering"],
            "tags": ["DevOps", "Kubernetes", "Docker", "CI/CD"]
        }
    ]
    
    # Filter jobs based on keyword if provided
    if keyword:
        keyword_lower = keyword.lower()
        filtered_jobs = [
            job for job in all_jobs 
            if keyword_lower in job['title'].lower() or 
               keyword_lower in job['company'].lower() or
               any(keyword_lower in tag.lower() for tag in job['tags'])
        ]
    else:
        filtered_jobs = all_jobs
    
    # Pagination logic
    jobs_per_page = 10
    start_idx = page * jobs_per_page
    end_idx = start_idx + jobs_per_page
    page_jobs = filtered_jobs[start_idx:end_idx]
    
    return {
        "page_count": (len(filtered_jobs) + jobs_per_page - 1) // jobs_per_page,
        "page": page,
        "jobs": page_jobs,
        "total_jobs": len(filtered_jobs),
        "source": "Mock Data"
    }

@app.get("/api/jobs")
async def search_jobs(
    page: int = 0,
    keyword: str = "software engineer",
    location: str = "United States",
    job_type: str = "full_time"
):
    """Search jobs using enhanced mock data (external APIs currently unavailable)"""
    print(f"Job search request: keyword='{keyword}', location='{location}', job_type='{job_type}', page={page}")
    
    # Use the enhanced mock data function
    try:
        return get_enhanced_mock_jobs(page, keyword, location)
    except Exception as e:
        print(f"Error with static data: {e}")
        import traceback
        traceback.print_exc()
        # Return the most basic response possible
        return {
            "page_count": 1,
            "page": 0,
            "jobs": [],
            "total_jobs": 0,
            "source": "Emergency Fallback"
        }

@app.get("/api/jobs/{job_id}")
async def get_job_details(job_id: str):
    """Get detailed information about a specific job"""
    try:
        job_id_int = int(job_id)
        
        # Get all jobs from the mock data function
        all_jobs_data = get_enhanced_mock_jobs(0, "", "")  # Get all jobs without filtering
        all_jobs = all_jobs_data.get("jobs", [])
        
        # Find job by ID
        job = next((j for j in all_jobs if j["id"] == job_id_int), None)
        
        if job:
            # Add additional details for the detailed view
            job["how_to_apply"] = "Please visit the company's careers page to apply for this position."
            job["company_url"] = job.get("landing_page", "#")
            return job
        else:
            raise HTTPException(status_code=404, detail="Job not found")
                
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid job ID")
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error fetching job details: {e}")
        raise HTTPException(status_code=500, detail="Failed to fetch job details")

